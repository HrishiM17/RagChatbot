import os
from typing import List, Dict, Any
import logging
from pathlib import Path
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from app.core.config import settings
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ": ", " ", ""]  # Better separators for context preservation
        )
    
    def read_pdf(self, file_path: str) -> str:
        """Read text from PDF file with improved error handling"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text += page_text + "\n\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1} of {file_path}: {e}")
                        continue
                
                # Clean up the extracted text
                text = self._clean_text(text)
                logger.info(f"Successfully extracted {len(text)} characters from PDF: {file_path}")
                return text
                
        except Exception as e:
            logger.error(f"Failed to read PDF {file_path}: {e}")
            return ""
    
    def read_docx(self, file_path: str) -> str:
        """Read text from DOCX file with improved formatting"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text += paragraph.text.strip() + "\n\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "\n"
            
            text = self._clean_text(text)
            logger.info(f"Successfully extracted {len(text)} characters from DOCX: {file_path}")
            return text
            
        except Exception as e:
            logger.error(f"Failed to read DOCX {file_path}: {e}")
            return ""
    
    def read_txt(self, file_path: str) -> str:
        """Read text from TXT file with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    text = self._clean_text(text)
                    logger.info(f"Successfully read {len(text)} characters from TXT file: {file_path} (encoding: {encoding})")
                    return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Failed to read TXT {file_path} with encoding {encoding}: {e}")
                continue
        
        logger.error(f"Failed to read TXT file {file_path} with any encoding")
        return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text for better processing"""
        if not text:
            return ""
        
        # Remove excessive whitespace while preserving structure
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Max 2 consecutive newlines
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'\n ', '\n', text)  # Remove spaces at start of lines
        
        # Fix common PDF extraction issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between lowercase and uppercase
        text = re.sub(r'([.!?])([A-Za-z])', r'\1 \2', text)  # Add space after punctuation
        
        # Remove page numbers and headers/footers (common patterns)
        text = re.sub(r'\n\d+\n', '\n', text)  # Standalone page numbers
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        
        return text.strip()
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file types with comprehensive error handling"""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return ""
        
        file_extension = Path(file_path).suffix.lower()
        file_size = os.path.getsize(file_path)
        
        logger.info(f"Processing file: {file_path} ({file_size} bytes, type: {file_extension})")
        
        # Check file size (skip very large files)
        max_file_size = 50 * 1024 * 1024  # 50MB limit
        if file_size > max_file_size:
            logger.warning(f"File too large ({file_size} bytes), skipping: {file_path}")
            return ""
        
        # Extract text based on file type
        text = ""
        if file_extension == '.pdf':
            text = self.read_pdf(file_path)
        elif file_extension == '.docx':
            text = self.read_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            text = self.read_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_extension} for file: {file_path}")
            return ""
        
        # Validate extracted text
        if not text or len(text.strip()) < 50:
            logger.warning(f"Little to no text extracted from {file_path} (length: {len(text)})")
            return ""
        
        return text
    
    def split_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """Split text into chunks with enhanced metadata"""
        if not text or not text.strip():
            return []
        
        try:
            chunks = self.text_splitter.split_text(text)
            
            if not chunks:
                logger.warning("Text splitter produced no chunks")
                return []
            
            result = []
            for i, chunk in enumerate(chunks):
                if len(chunk.strip()) < 20:  # Skip very small chunks
                    continue
                
                chunk_metadata = metadata.copy() if metadata else {}
                chunk_metadata.update({
                    "chunk_id": i,
                    "chunk_size": len(chunk),
                    "total_chunks": len(chunks),
                    "chunk_position": i / len(chunks),  # Relative position in document
                    "created_at": str(datetime.now())
                })
                
                result.append({
                    "text": chunk.strip(),
                    "metadata": chunk_metadata
                })
            
            logger.info(f"Successfully created {len(result)} chunks from {len(text)} characters")
            return result
            
        except Exception as e:
            logger.error(f"Failed to split text into chunks: {e}")
            return []
    
    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Process all supported files in a directory with comprehensive logging"""
        if not os.path.exists(directory_path):
            logger.error(f"Directory not found: {directory_path}")
            return []
        
        all_chunks = []
        supported_extensions = ['.pdf', '.docx', '.txt', '.md']
        processed_files = 0
        failed_files = 0
        
        logger.info(f"Starting directory processing: {directory_path}")
        
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_extension = Path(file).suffix.lower()
                
                if file_extension in supported_extensions:
                    logger.info(f"Processing file: {file}")
                    
                    # Extract text
                    text = self.extract_text_from_file(file_path)
                    
                    if text:
                        # Prepare comprehensive metadata
                        metadata = {
                            "source": file,
                            "file_path": file_path,
                            "file_type": file_extension,
                            "file_size": os.path.getsize(file_path),
                            "processed_at": str(datetime.now()),
                            "text_length": len(text)
                        }
                        
                        # Split into chunks
                        chunks = self.split_text(text, metadata)
                        
                        if chunks:
                            all_chunks.extend(chunks)
                            processed_files += 1
                            logger.info(f"✅ Successfully processed {file}: {len(chunks)} chunks created")
                        else:
                            failed_files += 1
                            logger.warning(f"❌ No chunks created from {file}")
                    else:
                        failed_files += 1
                        logger.warning(f"❌ No text extracted from {file}")
                else:
                    logger.debug(f"Skipping unsupported file: {file}")
        
        logger.info(f"Directory processing complete: {processed_files} files processed, {failed_files} files failed, {len(all_chunks)} total chunks created")
        return all_chunks
    
    def process_text_directly(self, text: str, source_name: str = "direct_input") -> List[Dict[str, Any]]:
        """Process text directly without file I/O"""
        if not text or not text.strip():
            logger.warning("Empty text provided for direct processing")
            return []
        
        # Clean the input text
        cleaned_text = self._clean_text(text)
        
        metadata = {
            "source": source_name,
            "file_type": "text",
            "input_method": "direct",
            "processed_at": str(datetime.now()),
            "text_length": len(cleaned_text)
        }
        
        chunks = self.split_text(cleaned_text, metadata)
        logger.info(f"Direct text processing: {len(chunks)} chunks created from {len(cleaned_text)} characters")
        return chunks
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get detailed information about a file"""
        try:
            if not os.path.exists(file_path):
                return {"error": "File not found"}
            
            file_stats = os.stat(file_path)
            file_extension = Path(file_path).suffix.lower()
            
            info = {
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "file_extension": file_extension,
                "file_size": file_stats.st_size,
                "file_size_mb": round(file_stats.st_size / (1024 * 1024), 2),
                "modified_time": str(datetime.fromtimestamp(file_stats.st_mtime)),
                "is_supported": file_extension in ['.pdf', '.docx', '.txt', '.md']
            }
            
            return info
            
        except Exception as e:
            logger.error(f"Failed to get file info for {file_path}: {e}")
            return {"error": str(e)}

# Import datetime for metadata
from datetime import datetime

# Global document processor instance
document_processor = DocumentProcessor()